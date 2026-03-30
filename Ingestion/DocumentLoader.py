from pathlib import Path
import logging
from typing import List, Optional
from Data_Structure import Document
import uuid

logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    Wrapper di alto livello per il caricamento: usato da IngestionPipeline.
    """
    
    def __init__(self, path):
        self.path = path
        self.loader = DocumentParser()

    def load_files(self)-> List[Document]:
        logger.info(f"Avvio caricamento documenti da: {self.path}")
        documents = self.loader.load(self.path)
        logger.info(f"Caricati {len(documents)} documenti totali")
        return documents
    


class DocumentParser:
    """
    Carica documenti da file PDF, TXT, Markdown e DOCX.
    Restituisce una lista di Document con metadati arricchiti.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
    
    def load(self, path: str) -> List[Document]:
        """
        Carica da un path: se è una directory carica tutti i file supportati,
        se è un singolo file carica solo quello.
        """
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"Path non trovato: {path}")
 
        if p.is_dir():
            return self._load_directory(p)
        else:
            doc = self._parse_file(p)
            return [doc] if doc else []
        
    def _load_directory(self, dir_path: Path) -> List[Document]:
        documents = []
        files = [
            f for f in dir_path.rglob("*")
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]
        logger.info(f"Trovati {len(files)} file in {dir_path}")
        for file_path in files:
            doc = self.load(file_path)
            if doc:
                documents.append(doc)
        return documents
    
    def _parse_file(self, file_path)-> Optional[Document]:
        ext = Path(file_path).suffix.lower()

        try:
            if ext in {".txt", ".md"}:
                content = self._parse_txt(file_path)
            elif ext == ".pdf":
                content = self._parse_pdf(file_path)
            elif ext == ".docx":
                content = self._parse_docx(file_path)
            else:
                logger.warning(f"Estensione non supportata: {file_path}")
                return None
            
            if not content or not content.strip():
                logger.warning(f"File vuoto o non leggibile: {file_path}")
                return None
 
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "extension": ext,
                "file_size_bytes": file_path.stat().st_size,
                "doc_id": str(uuid.uuid4()),
            }
            logger.info(f"Caricato: {file_path.name} ({len(content)} caratteri)")
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Errore durante il parsing del file {file_path}: {e}")
            return None
        
    def _parse_txt(self, path):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _parse_pdf(self, path):
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Installa pypdf: pip install pypdf")
        
        reader = PdfReader(str(path))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(f"[Pagina {i+1}]\n{text}")
        return "\n\n".join(pages_text)


    def _parse_docx(self, path):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Installa python-docx: pip install python-docx")

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        return "\n\n".join(paragraphs)
